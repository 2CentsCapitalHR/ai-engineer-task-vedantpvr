# docx_utils.py
from docx import Document
from docx.shared import Pt
from adgm_checklist import check_jurisdiction_paragraph, detect_missing_signature, detect_ambiguous_language
import re
import os
import copy
import json

def extract_docx_paragraphs(path):
    """
    Return a list of paragraph texts from .docx
    """
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip() != ""]
    return paragraphs

def summarize_issues(paragraphs, filename="docx", check_adgm=True):
    """
    Apply simple rules across paragraphs and return list of issues with metadata:
    [
        {
            "document": filename,
            "section": "Paragraph #n or heading",
            "issue": "...",
            "severity": "High",
            "suggestion": "...",
            "citation_key": "companies_regulation_jurisdiction" (optional)
        },
    ]
    """
    issues = []
    for idx, p in enumerate(paragraphs):
        if len(p.strip()) == 0:
            continue

        # Jurisdiction checks
        j_issues = check_jurisdiction_paragraph(p)
        for ji in j_issues:
            rec = {
                "document": filename,
                "section": f"Paragraph {idx+1}",
                "issue": ji["issue"],
                "severity": ji.get("severity", "Medium"),
                "suggestion": ji.get("suggestion"),
                "citation_key": ji.get("citation_key"),
            }
            issues.append(rec)

        # Ambiguity
        a_issues = detect_ambiguous_language(p)
        for ai in a_issues:
            rec = {
                "document": filename,
                "section": f"Paragraph {idx+1}",
                "issue": ai["issue"],
                "severity": ai.get("severity", "Low"),
                "suggestion": ai.get("suggestion"),
            }
            issues.append(rec)

    # Missing signature heuristic
    ms = detect_missing_signature(paragraphs)
    for m in ms:
        rec = {
            "document": filename,
            "section": "End of Document",
            "issue": m["issue"],
            "severity": m.get("severity", "High"),
            "suggestion": m.get("suggestion"),
        }
        issues.append(rec)

    # Additional checks can be added here (invalid clauses, incorrect formatting, etc.)
    return issues

def create_reviewed_docx(original_path, reviewed_path, issues):
    """
    Create a reviewed docx by copying original content and appending comment paragraphs
    near positions where issues were detected. Because python-docx doesn't support native
    Word comments, we append a clearly-styled paragraph after the flagged paragraph with marker "COMMENT".
    """
    src = Document(original_path)
    out = Document()
    # copy all elements while preserving text (styles preserved minimally)
    for p in src.paragraphs:
        newp = out.add_paragraph(p.text)
        # Basic style copying can be extended (font, bold, etc.)
    # Append a summary page with issues
    out.add_page_break()
    out.add_heading("Automated Review Comments / Issues", level=1)
    for idx, iss in enumerate(issues, 1):
        para = out.add_paragraph()
        para.add_run(f"[{idx}] Document: {iss.get('document')} | Section: {iss.get('section')}").bold = True
        out.add_paragraph(f"Issue: {iss.get('issue')}")
        if iss.get("suggestion"):
            s = out.add_paragraph(f"Suggestion: {iss.get('suggestion')}")
            s.style = out.styles['Intense Quote'] if 'Intense Quote' in out.styles else s.style
        if iss.get("citation_key"):
            c = out.add_paragraph(f"Citation: {iss.get('citation_key')}")
    out.save(reviewed_path)

def summarise_for_json(issues):
    """
    Return JSON-serializable summary for UI and download
    """
    return issues
